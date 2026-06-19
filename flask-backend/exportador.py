from docxtpl import DocxTemplate, Subdoc
import io
from xhtml2pdf import pisa
import os
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches

def set_table_borders(table):
    tbl = table._element
    tblPr = tbl.xpath('w:tblPr')[0]
    tblBorders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), '000000')
        tblBorders.append(border)
    tblPr.append(tblBorders)

def preparar_contexto_sesiones(sesiones_data, plantilla_extendida, doc=None):
    sesiones_context = []
    from datetime import datetime
    for s_data in sesiones_data:
        fecha_str = s_data.get('fecha_hora_inicio', '')
        duracion = ''
        fecha_rango = ''
        if fecha_str:
            if ' ' in fecha_str:
                # Calcular duracion si final esta presente
                final_str = s_data.get('fecha_hora_final', '')
                if final_str and ' ' in final_str:
                    try:
                        fmt = "%Y-%m-%d %H:%M:%S"
                        # Handle potential fractional seconds
                        inicio_dt = datetime.strptime(fecha_str.split('.')[0], fmt)
                        final_dt = datetime.strptime(final_str.split('.')[0], fmt)
                        mins = int((final_dt - inicio_dt).total_seconds() / 60)
                        duracion = f"{mins} min"
                        
                        f_inicio = inicio_dt.strftime("%d/%m/%Y %H:%M")
                        f_final = final_dt.strftime("%d/%m/%Y %H:%M")
                        fecha_rango = f"{f_inicio} - {f_final}"
                    except:
                        fecha_rango = fecha_str
                else:
                    fecha_rango = fecha_str
                fecha_str = fecha_str.split(' ')[0]
            else:
                fecha_rango = fecha_str
            
        obj_ids = s_data.get('objetivos_ids') or []
        tec_ids = s_data.get('tecnicas_ids') or []
        mat_ids = s_data.get('materiales_ids') or []
        cli_ids = s_data.get('clima_grupal_ids') or []
        log_ids = s_data.get('logros_ids') or []
        
        stars = s_data.get('cantidad_estrellas', 0)
        
        ctx = {
            'plantilla_extendida': plantilla_extendida,
            'nombre_sesion': s_data.get('nombre', 'Sesión sin nombre'),
            'tipo': 'Individual' if s_data.get('tipo', False) else 'Grupal',
            'modo': 'Presencial' if s_data.get('modo', False) else 'Virtual',
            
            'canciones_list': s_data.get('canciones_data', []),
            'emociones_list': s_data.get('emociones_data', []),
            
            'institucion_educativa': s_data.get('institucion_educativa', ''),
            'grado_seccion': s_data.get('grado_seccion', ''),
            'fecha': fecha_str,
            'fecha_rango': fecha_rango,
            'duracion': duracion,
            'facilitador': s_data.get('facilitador', ''),
            'numero_sesion': s_data.get('numero_sesion', ''),
            'numero_estudiantes': s_data.get('numero_estudiantes', ''),
            
            'part_alta': stars == 4,
            'part_media': stars == 3,
            'part_baja': stars == 2,
            'part_inconsistente': stars == 1,
            'participacion_obs': s_data.get('observaciones', ''),
            
            'individual': s_data.get('tipo', False),
            'presencial': s_data.get('modo', False),
            
            'obj_1': 1 in obj_ids,
            'obj_2': 2 in obj_ids,
            'obj_3': 3 in obj_ids,
            'obj_4': 4 in obj_ids,
            'obj_5': 5 in obj_ids,
            'obj_6': 6 in obj_ids,
            'obj_7': 7 in obj_ids,
            'obj_8': 8 in obj_ids,
            'obj_custom': s_data.get('objetivos_custom', ''),
            
            'tec_1': 1 in tec_ids,
            'tec_2': 2 in tec_ids,
            'tec_3': 3 in tec_ids,
            'tec_4': 4 in tec_ids,
            'tec_5': 5 in tec_ids,
            'tec_6': 6 in tec_ids,
            'tec_7': 7 in tec_ids,
            'tec_8': 8 in tec_ids,
            'tec_custom': s_data.get('tecnicas_custom', ''),
            
            'mat_1': 1 in mat_ids,
            'mat_2': 2 in mat_ids,
            'mat_3': 3 in mat_ids,
            'mat_4': 4 in mat_ids,
            'mat_5': 5 in mat_ids,
            'mat_6': 6 in mat_ids,
            'mat_custom': s_data.get('materiales_custom', ''),
            
            'inicio': s_data.get('inicio', ''),
            'actividad_central': s_data.get('actividad_central', ''),
            'cierre': s_data.get('cierre', ''),
            
            'cli_1': 1 in cli_ids,
            'cli_2': 2 in cli_ids,
            'cli_3': 3 in cli_ids,
            'cli_4': 4 in cli_ids,
            'cli_5': 5 in cli_ids,
            'cli_6': 6 in cli_ids,
            'cli_custom': s_data.get('clima_grupal_custom', ''),
            'cli_desc': s_data.get('descripcion_clima', ''),
            
            'log_1': 1 in log_ids,
            'log_2': 2 in log_ids,
            'log_3': 3 in log_ids,
            'log_4': 4 in log_ids,
            'log_5': 5 in log_ids,
            'log_6': 6 in log_ids,
            
            'dificultades': s_data.get('dificultades', ''),
            'recomendaciones': s_data.get('recomendaciones', ''),
        }
        
        if plantilla_extendida:
            # We will use unique placeholders and replace them manually after render
            # because Subdoc rendering in lists can be flaky.
            ctx['canciones_table'] = f"[[TABLE_SONGS_{len(sesiones_context)}]]"
            ctx['emociones_table'] = f"[[TABLE_EMOCIONES_{len(sesiones_context)}]]"
            
            # Pre-generate tables to be used later
            from docx import Document
            temp_doc = Document()
            
            # Songs Table
            if s_data.get('canciones_data'):
                table = temp_doc.add_table(rows=1, cols=3)
                table.autofit = False
                set_table_borders(table)
                hdr_cells = table.rows[0].cells
                hdr_cells[0].text = 'Título'
                hdr_cells[1].text = 'Autor'
                hdr_cells[2].text = 'Álbum'
                widths = [Inches(2.5), Inches(2.0), Inches(1.5)]
                for i, width in enumerate(widths):
                    table.columns[i].width = width
                for c in s_data.get('canciones_data', []):
                    row_cells = table.add_row().cells
                    row_cells[0].text = str(c.get('nombre', ''))
                    row_cells[1].text = str(c.get('autor', ''))
                    row_cells[2].text = str(c.get('album', ''))
            else:
                table = temp_doc.add_paragraph('Ninguna canción seleccionada.')
            
            # Emotions Table
            temp_doc_emo = Document()
            if s_data.get('emociones_data'):
                table_emo = temp_doc_emo.add_table(rows=1, cols=2)
                table_emo.autofit = False
                set_table_borders(table_emo)
                hdr_emo = table_emo.rows[0].cells
                hdr_emo[0].text = 'Emoción (Plutchik)'
                hdr_emo[1].text = 'Palabra seleccionada'
                widths_emo = [Inches(3.0), Inches(3.0)]
                for i, width in enumerate(widths_emo):
                    table_emo.columns[i].width = width
                for e in s_data.get('emociones_data', []):
                    row_cells = table_emo.add_row().cells
                    row_cells[0].text = str(e.get('emocion', ''))
                    row_cells[1].text = str(e.get('palabra', ''))
            else:
                table_emo = temp_doc_emo.add_paragraph('Ninguna emoción seleccionada.')
                
            # Store them in a special place to find them after render
            
            ctx['_table_songs_obj'] = table
            ctx['_table_emo_obj'] = table_emo
            
        sesiones_context.append(ctx)
    return {'sesiones': sesiones_context}

def crear_docx_sesiones(sesiones_data, plantilla_extendida):
    template_path = os.path.join(os.path.dirname(__file__), 'FICHA_JINJA_TEMPLATE.docx')
    if not os.path.exists(template_path):
        print(f"Error: Plantilla no encontrada en {template_path}")
        return None
        
    try:
        doc = DocxTemplate(template_path)
        context = preparar_contexto_sesiones(sesiones_data, plantilla_extendida, doc)
        doc.render(context)
        
        # Manual injection of tables
        for s in context['sesiones']:
            if '_table_songs_obj' in s:
                placeholder_songs = s['canciones_table']
                placeholder_emo = s['emociones_table']
                
                # Find paragraphs with placeholders
                for p in doc.paragraphs:
                    if placeholder_songs in p.text:
                        # Replace paragraph with table
                        table_xml = s['_table_songs_obj']
                        if hasattr(table_xml, '_element'): # It's a table
                            p._p.addnext(table_xml._element)
                        else: # It's a paragraph
                            p._p.addnext(table_xml._p)
                        p._p.getparent().remove(p._p)
                    
                    if placeholder_emo in p.text:
                        table_xml = s['_table_emo_obj']
                        if hasattr(table_xml, '_element'):
                            p._p.addnext(table_xml._element)
                        else:
                            p._p.addnext(table_xml._p)
                        p._p.getparent().remove(p._p)
        
        file_stream = io.BytesIO()
        doc.save(file_stream)
        file_stream.seek(0)
        return file_stream
    except Exception as e:
        print(f"Error al generar DOCX con docxtpl: {e}")
        return None

def crear_pdf_sesiones(sesiones_data, plantilla_extendida):
    context = preparar_contexto_sesiones(sesiones_data, plantilla_extendida)
    
    html = """
    <html>
    <head>
    <style>
        body { font-family: Helvetica, sans-serif; font-size: 11pt; }
        h1 { text-align: center; font-size: 14pt; margin-bottom: 20px; }
        .section-title { font-weight: bold; font-size: 12pt; margin-top: 15px; margin-bottom: 5px; }
        .grid { display: block; margin-bottom: 10px; }
        .row { margin-bottom: 5px; }
        .label { font-weight: bold; }
        .box { font-family: DejaVu Sans, sans-serif; font-size: 14pt; }
        .textarea { border: 1px solid #ccc; min-height: 50px; padding: 5px; width: 100%; margin-top: 5px; }
    </style>
    </head>
    <body>
    """
    
    for i, s in enumerate(context['sesiones']):
        if i > 0:
            html += "<pdf:nextpage />"
            
        check = lambda val: "&#9745;" if val else "&#9744;"
        
        html += f"""
        <h1>FICHA DE REGISTRO<br>MÚSICA COMO TERAPIA</h1>
        
        <div class="section-title">Datos generales</div>
        <div class="row"><span class="label">Nombre de sesión:</span> {s['nombre_sesion']}</div>
        <div class="row"><span class="label">Tipo:</span> {s['tipo']} &nbsp;&nbsp; <span class="label">Modo:</span> {s['modo']}</div>
        <div class="row"><span class="label">Institución educativa:</span> {s['institucion_educativa']}</div>
        <div class="row"><span class="label">Grado / Sección:</span> {s['grado_seccion']}</div>
        <div class="row"><span class="label">Número de estudiantes:</span> {s['numero_estudiantes']}</div>
        <div class="row"><span class="label">Fecha:</span> {s['fecha']}</div>
        <div class="row"><span class="label">Duración de la sesión:</span> {s['duracion']}</div>
        <div class="row"><span class="label">N° de sesión:</span> {s['numero_sesion']}</div>
        <div class="row"><span class="label">Facilitador(a):</span> {s['facilitador']}</div>
        """
        
        if s.get('plantilla_extendida', False):
            # Canciones Table
            html += '<div class="section-title">Canciones seleccionadas</div>'
            html += '<table border="1" width="100%" cellpadding="3" cellspacing="0"><tr><th>Título</th><th>Autor</th><th>Álbum</th></tr>'
            for c in s['canciones_list']:
                html += f"<tr><td>{c.get('nombre','')}</td><td>{c.get('autor','')}</td><td>{c.get('album','')}</td></tr>"
            html += "</table>"
            
            # Emociones Table
            html += '<div class="section-title">Emociones (Rueda)</div>'
            html += '<table border="1" width="100%" cellpadding="3" cellspacing="0"><tr><th>Emoción (Plutchik)</th><th>Palabra seleccionada</th></tr>'
            for e in s['emociones_list']:
                html += f"<tr><td>{e.get('emocion','')}</td><td>{e.get('palabra','')}</td></tr>"
            html += "</table>"
            
        html += f"""
        <br/>
        <div class="row"><span class="label">Tipo de sesión:</span> <span class="box">{check(s['individual'])}</span> Individual &nbsp;&nbsp; <span class="box">{check(not s['individual'])}</span> Grupal</div>
        <div class="row"><span class="label">Modalidad:</span> <span class="box">{check(s['presencial'])}</span> Presencial &nbsp;&nbsp; <span class="box">{check(not s['presencial'])}</span> Virtual</div>
        
        <div class="section-title">Objetivos de la sesión</div>
        <div><span class="box">{check(s['obj_1'])}</span> Cohesión grupal – trabajo colaborativo</div>
        <div><span class="box">{check(s['obj_2'])}</span> Regulación emocional colectiva</div>
        <div><span class="box">{check(s['obj_3'])}</span> Mejora de la convivencia y respeto</div>
        <div><span class="box">{check(s['obj_4'])}</span> Estimulación de la atención grupal</div>
        <div><span class="box">{check(s['obj_5'])}</span> Expresión emocional y creativa</div>
        <div><span class="box">{check(s['obj_6'])}</span> Comunicación verbal y no verbal</div>
        <div><span class="box">{check(s['obj_7'])}</span> Desarrollo rítmico y coordinación motora</div>
        <div><span class="box">{check(s['obj_8'])}</span> Disminución de ansiedad o tensión en el aula</div>
        <div>Otros: {s['obj_custom']}</div>
        
        <div class="section-title">Técnicas utilizadas</div>
        <div><span class="box">{check(s['tec_1'])}</span> Canto grupal de canciones dirigidas</div>
        <div><span class="box">{check(s['tec_2'])}</span> Improvisación musical grupal (instrumental o vocal)</div>
        <div><span class="box">{check(s['tec_3'])}</span> Percusión corporal</div>
        <div><span class="box">{check(s['tec_4'])}</span> Movimiento creativo con música</div>
        <div><span class="box">{check(s['tec_5'])}</span> Audición receptiva</div>
        <div><span class="box">{check(s['tec_6'])}</span> Creación colectiva de canciones / letras</div>
        <div><span class="box">{check(s['tec_7'])}</span> Actividades rítmicas en círculo</div>
        <div><span class="box">{check(s['tec_8'])}</span> Dinámicas de coordinación y seguimiento</div>
        <div>Otras: {s['tec_custom']}</div>
        
        <div class="section-title">Materiales utilizados</div>
        <div><span class="box">{check(s['mat_1'])}</span> Instrumentos de percusión menor</div>
        <div><span class="box">{check(s['mat_2'])}</span> Cajones / tambores</div>
        <div><span class="box">{check(s['mat_3'])}</span> Parlantes / música grabada</div>
        <div><span class="box">{check(s['mat_4'])}</span> Carteles con letras</div>
        <div><span class="box">{check(s['mat_5'])}</span> Objetos sonoros</div>
        <div><span class="box">{check(s['mat_6'])}</span> Aros / telas / cintas rítmicas</div>
        <div>Otros: {s['mat_custom']}</div>
        
        <div class="section-title">Desarrollo de la sesión</div>
        <div class="label">Inicio (calentamiento – vínculo):</div>
        <div class="textarea">{s['inicio'].replace(chr(10), '<br>')}</div>
        <div class="label">Actividad central (trabajo musical):</div>
        <div class="textarea">{s['actividad_central'].replace(chr(10), '<br>')}</div>
        <div class="label">Cierre (relajación – reflexión – retroalimentación):</div>
        <div class="textarea">{s['cierre'].replace(chr(10), '<br>')}</div>
        
        <div class="section-title">Observaciones del grupo</div>
        <div class="label">Clima grupal:</div>
        <div><span class="box">{check(s['cli_1'])}</span> Armónico</div>
        <div><span class="box">{check(s['cli_2'])}</span> Disperso</div>
        <div><span class="box">{check(s['cli_3'])}</span> Agitado</div>
        <div><span class="box">{check(s['cli_4'])}</span> Colaborativo</div>
        <div><span class="box">{check(s['cli_5'])}</span> Con conflictos Descripción: {s['cli_custom']}</div>
        
        <div class="label">Participación:</div>
        <div><span class="box">{check(s['part_alta'])}</span> Alta</div>
        <div><span class="box">{check(s['part_media'])}</span> Media</div>
        <div><span class="box">{check(s['part_baja'])}</span> Baja</div>
        <div><span class="box">{check(s['part_inconsistente'])}</span> Inconsistente Observaciones: {s['participacion_obs']}</div>
        
        <div class="section-title">Logros observados en el grupo</div>
        <div><span class="box">{check(s['log_1'])}</span> Mejoró el clima emocional del aula</div>
        <div><span class="box">{check(s['log_2'])}</span> Hubo mayor cohesión y cooperación</div>
        <div><span class="box">{check(s['log_3'])}</span> Aumentó el nivel de energía positiva</div>
        <div><span class="box">{check(s['log_4'])}</span> Se redujo la tensión o conflicto</div>
        <div><span class="box">{check(s['log_5'])}</span> Incrementó la atención y seguimiento de instrucciones</div>
        <div><span class="box">{check(s['log_6'])}</span> Expresaron emociones a través de la música</div>
        
        <div class="section-title">Dificultades o necesidades del grupo</div>
        <div class="textarea">{s['dificultades'].replace(chr(10), '<br>')}</div>
        
        <div class="section-title">Recomendaciones para próximas sesiones</div>
        <div class="textarea">{s['recomendaciones'].replace(chr(10), '<br>')}</div>
        """
        
    html += "</body></html>"
    
    file_stream = io.BytesIO()
    # PISA requires utf-8 encode
    pisa_status = pisa.CreatePDF(html.encode('utf-8'), dest=file_stream)
    
    if pisa_status.err:
        print("Error generando PDF:", pisa_status.err)
        return None
        
    file_stream.seek(0)
    return file_stream
