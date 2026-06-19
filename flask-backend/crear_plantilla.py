import docx
import re
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def set_table_borders(table):
    tbl = table._element
    tblPr = tbl.xpath('w:tblPr')[0]
    tblBorders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4') # border size
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), '000000') # black
        tblBorders.append(border)
    tblPr.append(tblBorders)

def create_jinja_template():
    doc = docx.Document('FICHA DE REGISTRO.docx')
    
    # We want to insert {% for s in sesiones %} at the very beginning
    # and {% endfor %} at the very end.
    
    # Text replacements map
    replacements = {
        "Institución educativa:": "Institución educativa: {{ s.institucion_educativa }}",
        "Grado / Sección:": "Grado / Sección: {{ s.grado_seccion }}",
        "Número de estudiantes:": "Número de estudiantes: {{ s.numero_estudiantes }}",
        "Duración de la sesión:": "Duración de la sesión: {{ s.duracion }}",
        "N° de sesión:": "N° de sesión: {{ s.numero_sesion }}",
        "Facilitador(a):": "Facilitador(a): {{ s.facilitador }}\nNombre de sesión: {{ s.nombre_sesion }}\nTipo: {{ s.tipo }}\nModo: {{ s.modo }}",
        
        "☐ Individual": "{% if s.individual %}☑{% else %}☐{% endif %} Individual",
        "☐ Grupal": "{% if not s.individual %}☑{% else %}☐{% endif %} Grupal",
        "☐ Presencial": "{% if s.presencial %}☑{% else %}☐{% endif %} Presencial",
        "☐ Virtual": "{% if not s.presencial %}☑{% else %}☐{% endif %} Virtual",
        
        # Objetivos
        "☐ Cohesión grupal – trabajo colaborativo": "{% if s.obj_1 %}☑{% else %}☐{% endif %} Cohesión grupal – trabajo colaborativo",
        "☐ Regulación emocional colectiva": "{% if s.obj_2 %}☑{% else %}☐{% endif %} Regulación emocional colectiva",
        "☐ Mejora de la convivencia y respeto": "{% if s.obj_3 %}☑{% else %}☐{% endif %} Mejora de la convivencia y respeto",
        "☐ Estimulación de la atención grupal": "{% if s.obj_4 %}☑{% else %}☐{% endif %} Estimulación de la atención grupal",
        "☐ Expresión emocional y creativa": "{% if s.obj_5 %}☑{% else %}☐{% endif %} Expresión emocional y creativa",
        "☐ Comunicación verbal y no verbal": "{% if s.obj_6 %}☑{% else %}☐{% endif %} Comunicación verbal y no verbal",
        "☐ Desarrollo rítmico y coordinación motora": "{% if s.obj_7 %}☑{% else %}☐{% endif %} Desarrollo rítmico y coordinación motora",
        "☐ Disminución de ansiedad o tensión en el aula": "{% if s.obj_8 %}☑{% else %}☐{% endif %} Disminución de ansiedad o tensión en el aula",
        "Otros:": "Otros: {{ s.obj_custom }}",
        
        # Tecnicas
        "☐ Canto grupal de canciones dirigidas": "{% if s.tec_1 %}☑{% else %}☐{% endif %} Canto grupal de canciones dirigidas",
        "☐ Improvisación musical grupal (instrumental o vocal)": "{% if s.tec_2 %}☑{% else %}☐{% endif %} Improvisación musical grupal (instrumental o vocal)",
        "☐ Percusión corporal": "{% if s.tec_3 %}☑{% else %}☐{% endif %} Percusión corporal",
        "☐ Movimiento creativo con música": "{% if s.tec_4 %}☑{% else %}☐{% endif %} Movimiento creativo con música",
        "☐ Audición receptiva": "{% if s.tec_5 %}☑{% else %}☐{% endif %} Audición receptiva",
        "☐ Creación colectiva de canciones / letras": "{% if s.tec_6 %}☑{% else %}☐{% endif %} Creación colectiva de canciones / letras",
        "☐ Actividades rítmicas en círculo": "{% if s.tec_7 %}☑{% else %}☐{% endif %} Actividades rítmicas en círculo",
        "☐ Dinámicas de coordinación y seguimiento": "{% if s.tec_8 %}☑{% else %}☐{% endif %} Dinámicas de coordinación y seguimiento",
        "Otras:": "Otras: {{ s.tec_custom }}",
        
        # Materiales
        "☐ Instrumentos de percusión menor": "{% if s.mat_1 %}☑{% else %}☐{% endif %} Instrumentos de percusión menor",
        "☐ Cajones / tambores": "{% if s.mat_2 %}☑{% else %}☐{% endif %} Cajones / tambores",
        "☐ Parlantes / música grabada": "{% if s.mat_3 %}☑{% else %}☐{% endif %} Parlantes / música grabada",
        "☐ Carteles con letras": "{% if s.mat_4 %}☑{% else %}☐{% endif %} Carteles con letras",
        "☐ Objetos sonoros": "{% if s.mat_5 %}☑{% else %}☐{% endif %} Objetos sonoros",
        "☐ Aros / telas / cintas rítmicas": "{% if s.mat_6 %}☑{% else %}☐{% endif %} Aros / telas / cintas rítmicas",
        "Otros:": "Otros: {{ s.mat_custom }}",
        
        # Desarrollo
        "Inicio (calentamiento – vínculo):": "Inicio (calentamiento – vínculo):\n{{ s.inicio }}",
        "Actividad central (trabajo musical):": "Actividad central (trabajo musical):\n{{ s.actividad_central }}",
        "Cierre (relajación – reflexión – retroalimentación):": "Cierre (relajación – reflexión – retroalimentación):\n{{ s.cierre }}",
        
        # Observaciones
        "☐ Armónico": "{% if s.cli_1 %}☑{% else %}☐{% endif %} Armónico",
        "☐ Disperso": "{% if s.cli_2 %}☑{% else %}☐{% endif %} Disperso",
        "☐ Agitado": "{% if s.cli_3 %}☑{% else %}☐{% endif %} Agitado",
        "☐ Colaborativo": "{% if s.cli_4 %}☑{% else %}☐{% endif %} Colaborativo",
        "☐ Con conflictos Descripción:": "{% if s.cli_5 %}☑{% else %}☐{% endif %} Con conflictos Descripción: {{ s.cli_custom }}",
        
        "☐ Alta": "{% if s.part_alta %}☑{% else %}☐{% endif %} Alta",
        "☐ Media": "{% if s.part_media %}☑{% else %}☐{% endif %} Media",
        "☐ Baja": "{% if s.part_baja %}☑{% else %}☐{% endif %} Baja",
        "☐ Inconsistente Observaciones:": "{% if s.part_inconsistente %}☑{% else %}☐{% endif %} Inconsistente Observaciones: {{ s.participacion_obs }}",
        
        # Logros
        "☐ Mejoró el clima emocional del aula": "{% if s.log_1 %}☑{% else %}☐{% endif %} Mejoró el clima emocional del aula",
        "☐ Hubo mayor cohesión y cooperación": "{% if s.log_2 %}☑{% else %}☐{% endif %} Hubo mayor cohesión y cooperación",
        "☐ Aumentó el nivel de energía positiva": "{% if s.log_3 %}☑{% else %}☐{% endif %} Aumentó el nivel de energía positiva",
        "☐ Se redujo la tensión o conflicto": "{% if s.log_4 %}☑{% else %}☐{% endif %} Se redujo la tensión o conflicto",
        "☐ Incrementó la atención y seguimiento de instrucciones": "{% if s.log_5 %}☑{% else %}☐{% endif %} Incrementó la atención y seguimiento de instrucciones",
        "☐ Expresaron emociones a través de la música": "{% if s.log_6 %}☑{% else %}☐{% endif %} Expresaron emociones a través de la música",
        
        "Dificultades o necesidades del grupo": "Dificultades o necesidades del grupo\n{{ s.dificultades }}",
        "Recomendaciones para próximas sesiones": "Recomendaciones para próximas sesiones\n{{ s.recomendaciones }}"
    }
    
    # Helper to replace keeping basic formatting if possible, 
    # but since it's simple we can just replace paragraph.text.
    # Note: replacing paragraph.text clears runs, so any bolding inside the replaced part is lost.
    # The template has very basic styling so this might be acceptable.
    
    otros_targets = [
        "{{ s.obj_custom }}",
        "{{ s.tec_custom }}",
        "{{ s.mat_custom }}"
    ]
    otros_count = 0

    for p in doc.paragraphs:
        original = p.text
        # Special case for "Otros:" / "Otros: _____"
        if "Otros:" in original and otros_count < len(otros_targets):
            p.text = original.replace("Otros: _______________________________________________________", f"Otros: {otros_targets[otros_count]}")
            original = p.text
            otros_count += 1
            
        # Robust Date replacement
        if p.text.strip().startswith("Fecha:"):
            p.text = f"Fecha: {{{{ s.fecha_rango }}}}"
            original = p.text

        for k, v in replacements.items():
            if k in original:
                p.text = original.replace(k, v)
                original = p.text
                
    # Insert Placeholders for Extended Template
    # We find the Facilitador(a) paragraph to insert after
    for i, p in enumerate(doc.paragraphs):
        if "{{ s.facilitador }}" in p.text:
            p_ext_start = doc.paragraphs[i].insert_paragraph_before('{% if s.plantilla_extendida %}')
            
            p_songs_title = doc.paragraphs[i+1].insert_paragraph_before('Canciones seleccionadas:')
            p_songs_title.runs[0].bold = True
            p_songs_table = doc.add_paragraph('{{ s.canciones_table }}')
            p_songs_title._p.addnext(p_songs_table._p)
            
            p_emo_title = doc.add_paragraph('Emociones (Rueda):')
            p_emo_title.runs[0].bold = True
            p_emo_table = doc.add_paragraph('{{ s.emociones_table }}')
            p_songs_table._p.addnext(p_emo_title._p)
            p_emo_title._p.addnext(p_emo_table._p)
            
            p_ext_end = doc.add_paragraph('{% endif %}')
            p_emo_table._p.addnext(p_ext_end._p)
            break
                
    # Add Jinja loop tags
    doc.paragraphs[0].insert_paragraph_before('{% for s in sesiones %}')

    
    p_end = doc.add_paragraph('{% endfor %}')
    
    doc.save('FICHA_JINJA_TEMPLATE.docx')
    print("FICHA_JINJA_TEMPLATE.docx creada con exito.")

if __name__ == '__main__':
    create_jinja_template()
